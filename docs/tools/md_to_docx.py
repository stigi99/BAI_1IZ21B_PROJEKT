#!/usr/bin/env python3
"""Convert the BAI Markdown report to a readable DOCX.

This is intentionally small and local to the report. It handles the Markdown
features used in docs/Sprawozdanie_BAI.md without requiring pandoc.
"""

from __future__ import annotations

import argparse
import html
import os
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pygments import lex
from pygments.lexers import get_lexer_by_name, guess_lexer
from pygments.token import Comment, Keyword, Name, Number, Operator, String, Token


IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")

TOC_PAGE_NUMBERS = {
    "Streszczenie": 1,
    "Spis rysunków i tabel": 2,
    "0. Informacje organizacyjne i podział pracy": 4,
    "1. Cel projektu": 5,
    "2. Zakres funkcjonalny aplikacji": 6,
    "3. Architektura i uruchomienie": 7,
    "4. Mechanizm vulnerable / secure": 9,
    "5. Dokumentacja funkcjonalności użytkowej": 12,
    "6. Zestawienie podatności": 14,
    "6A. Infografiki, diagramy i wizualizacje AI": 15,
    "6B. Podatności jako funkcjonalności aplikacji": 26,
    "7. SQL Injection": 27,
    "8. Stored XSS": 31,
    "9. Broken Authentication": 35,
    "10. Broken Access Control / IDOR": 39,
    "11. CSRF": 43,
    "12. Sensitive Data Exposure": 46,
    "13. Path Traversal / LFI": 49,
    "14. Command Injection": 54,
    "15. Testy integracyjne": 58,
    "16. Propozycje dalszych zmian w kodzie": 59,
    "17. Wnioski końcowe": 62,
    "18. Załącznik: scenariusz demonstracji": 62,
    "19. Załącznik: mapa ekranów aplikacji": 65,
    "20. Załącznik: macierz ryzyka": 67,
    "21. Załącznik: analiza plików": 69,
    "22. Załącznik: szczegółowa checklista testowania manualnego": 70,
    "23. Załącznik: propozycja narracji na obronę": 72,
    "24. Załącznik: rekomendowany format PDF": 74,
    "25. Krótkie podsumowanie dla prowadzącego": 75,
    "26. Diagramy Mermaid do wersji elektronicznej": 75,
    "27. Realne incydenty, CVE i ciekawostki": 77,
    "28. Literatura, książki, artykuły naukowe i źródła branżowe": 80,
}

ACTIVE_TOC_PAGE_NUMBERS: dict[str, int] = {}
FIGURE_COUNTER = 0
LISTING_COUNTER = 0
TABLE_COUNTER = 0


def code_caption_title(code: list[str], language: str) -> str:
    raw = "\n".join(code)
    compact = " ".join(line.strip() for line in code if line.strip())
    lang = (language or "tekst").lower()

    if "sqlmap " in raw:
        return "Komenda sqlmap potwierdzająca podatność SQL Injection"
    if "curl " in raw:
        return "Żądanie curl użyte do weryfikacji endpointu HTTP"
    if "http --" in raw:
        return "Żądanie HTTPie użyte do weryfikacji endpointu HTTP"
    if "go test" in raw:
        return "Komendy uruchamiające testy automatyczne projektu"
    if "UNION SELECT" in raw:
        return "Payload SQL Injection użyty w wyszukiwarce biblioteki"
    if lang == "http":
        return "Żądanie HTTP odtwarzające scenariusz ataku"
    if lang == "sql":
        return "Fragment zapytania lub payloadu SQL"
    if lang in {"go", "templ", "html", "css", "javascript", "js"}:
        return f"Fragment implementacji aplikacji ({lang})"
    if compact:
        return f"Fragment materiału testowego ({lang})"
    return f"Kod {lang}"


def collect_report_items(lines: list[str]) -> dict[str, list[dict[str, str | int]]]:
    """Collect figures, tables and listings before rendering front matter."""
    items: dict[str, list[dict[str, str | int]]] = {"figures": [], "tables": [], "listings": []}
    in_code = False
    language = ""
    code_lines: list[str] = []
    index = 0
    figure_no = 0
    table_no = 0
    listing_no = 0

    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                listing_no += 1
                items["listings"].append(
                    {
                        "number": listing_no,
                        "title": code_caption_title(code_lines, language),
                        "anchor": f"lst{listing_no}",
                    }
                )
                code_lines = []
                language = ""
                in_code = False
            else:
                in_code = True
                language = line.strip("`").strip()
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        image_match = IMAGE_RE.search(line)
        if image_match:
            figure_no += 1
            items["figures"].append(
                {"number": figure_no, "title": image_match.group(1), "anchor": f"fig{figure_no}"}
            )
            index += 1
            continue

        if line.lstrip().startswith("|"):
            rows, next_index = parse_table(lines, index)
            if rows:
                table_no += 1
                title = clean_inline(rows[0][0]) if rows and rows[0] else "Zestawienie danych raportowych"
                items["tables"].append(
                    {"number": table_no, "title": title, "anchor": f"tbl{table_no}"}
                )
            index = next_index
            continue

        index += 1

    return items


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(clean_inline(text))
    run.bold = bold
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9.5)


def clean_inline(text: str) -> str:
    text = html.unescape(text.strip())
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = LINK_RE.sub(r"\1 (\2)", text)
    return text


def add_inline_runs(paragraph, text: str) -> None:
    """Add basic inline formatting for backticks and strong emphasis."""
    text = LINK_RE.sub(r"\1 (\2)", text)
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(html.unescape(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(html.unescape(token[1:-1]))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(html.unescape(token[2:-2]))
            run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(html.unescape(text[cursor:]))


def token_color(token_type) -> str:
    if token_type in Keyword or token_type in Name.Builtin:
        return "7C3AED"
    if token_type in String:
        return "B91C1C"
    if token_type in Comment:
        return "64748B"
    if token_type in Number:
        return "0E7490"
    if token_type in Operator:
        return "BE185D"
    if token_type in Name.Function or token_type in Name.Class:
        return "0369A1"
    if token_type in Name.Attribute or token_type in Name.Tag:
        return "047857"
    return "111827"


def code_source(language: str) -> str:
    lang = (language or "").lower()
    if lang in {"go", "templ", "html", "css", "js", "javascript"}:
        return "fragment kodu źródłowego projektu"
    if lang in {"sql", "http", "bash", "text"}:
        return "materiał testowy opracowany na potrzeby weryfikacji"
    return "opracowanie własne"


def add_code_block(document: Document, code: list[str], language: str) -> None:
    global LISTING_COUNTER
    LISTING_COUNTER += 1

    raw = "\n".join(code)
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CBD5E1")
        borders.append(border)
    tc_pr.append(borders)

    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles["No Spacing"]
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True

    try:
        lexer = get_lexer_by_name(language or "text")
    except Exception:
        try:
            lexer = guess_lexer(raw)
        except Exception:
            lexer = get_lexer_by_name("text")

    for token_type, value in lex(raw, lexer):
        if not value:
            continue
        parts = value.split("\n")
        for idx, part in enumerate(parts):
            if idx:
                paragraph.add_run().add_break()
            if part:
                run = paragraph.add_run(part)
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor.from_string(token_color(token_type))
                if token_type in Keyword:
                    run.bold = True

    caption = document.add_paragraph()
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False
    title = code_caption_title(code, language)
    run = caption.add_run(
        f"Listing {LISTING_COUNTER}. {title}. Źródło: {code_source(language)}."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("475569")
    add_heading_bookmark(caption, f"lst{LISTING_COUNTER}", 2000 + LISTING_COUNTER)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        raw = lines[index].strip()
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    global TABLE_COUNTER
    if not rows:
        return
    TABLE_COUNTER += 1
    caption = document.add_paragraph()
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_after = Pt(4)
    run = caption.add_run(f"Tabela {TABLE_COUNTER}. Zestawienie danych raportowych. Źródło: opracowanie własne.")
    run.bold = True
    run.font.size = Pt(9)
    add_heading_bookmark(caption, f"tbl{TABLE_COUNTER}", 3000 + TABLE_COUNTER)
    max_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for col_index in range(max_cols):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, text, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
    document.add_paragraph()


def convert_svg(svg_path: Path, assets_dir: Path) -> Path | None:
    png_path = assets_dir / f"{svg_path.stem}.png"
    if png_path.exists() and png_path.stat().st_mtime >= svg_path.stat().st_mtime:
        return png_path
    try:
        subprocess.run(
            ["rsvg-convert", "-w", "1600", "-o", str(png_path), str(svg_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        return png_path
    except Exception:
        return None


def image_width_inches(target: str) -> float:
    """Choose image width by asset class to avoid sparse A4 pages."""
    normalized = target.replace("\\", "/")
    if "generated/assets/ai/" in normalized:
        return 6.25
    if "generated/assets/crops/" in normalized:
        return 5.1
    if "screenshots/" in normalized:
        return 5.1
    if "generated/assets/mermaid/" in normalized:
        return 5.95
    return 6.15


def image_source(target: str) -> str:
    normalized = target.replace("\\", "/")
    if "generated/assets/ai/" in normalized:
        return "grafika wygenerowana AI na potrzeby projektu"
    if "generated/assets/mermaid/" in normalized:
        return "diagram Mermaid wygenerowany z opisu architektury"
    if "generated/assets/crops/" in normalized or "screenshots/" in normalized:
        return "zrzut ekranu z aplikacji MikuMiku Fan Hub"
    if "sqlmap" in normalized:
        return "wynik narzędzia sqlmap uruchomionego lokalnie"
    if "curl" in normalized:
        return "wynik narzędzia curl uruchomionego lokalnie"
    if "httpie" in normalized:
        return "wynik narzędzia HTTPie uruchomionego lokalnie"
    if "generated/assets/" in normalized:
        return "opracowanie własne"
    return "materiał własny projektu"


def add_image(document: Document, md_path: Path, output_dir: Path, alt: str, target: str) -> None:
    global FIGURE_COUNTER
    image_path = (md_path.parent / target).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Brak obrazu: {target}]").italic = True
        return

    usable_path = image_path
    if image_path.suffix.lower() == ".svg":
        converted = convert_svg(image_path, output_dir / "assets")
        if converted is None:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"[SVG: {alt} - {target}]").italic = True
            return
        usable_path = converted

    try:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        paragraph.alignment = 1
        paragraph.add_run().add_picture(str(usable_path), width=Inches(image_width_inches(target)))
        FIGURE_COUNTER += 1
        caption = document.add_paragraph()
        caption.alignment = 1
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(8)
        caption_run = caption.add_run(f"Rysunek {FIGURE_COUNTER}. {alt}. Źródło: {image_source(target)}.")
        caption_run.bold = True
        caption_run.font.size = Pt(9)
        caption_run.font.color.rgb = RGBColor.from_string("475569")
        add_heading_bookmark(caption, f"fig{FIGURE_COUNTER}", 1000 + FIGURE_COUNTER)
    except Exception:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Nie udało się osadzić obrazu: {target}]").italic = True


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_page_number_footer(document)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.28

    for level, size, color, space_before in [
        (1, 22, "17324D", 24),
        (2, 16, "1F4E79", 18),
        (3, 13, "305496", 11),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(5)
        # Keep a heading glued to the text that follows it so no heading is
        # ever stranded alone at the bottom of a page.
        style.paragraph_format.keep_with_next = True


def collect_section_titles(lines: list[str]) -> list[tuple[int, str, str]]:
    """Ordered numbered headings for the table of contents."""
    titles: list[tuple[int, str, str]] = []
    in_code = False
    h2 = 0
    h3 = 0
    for line in lines:
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if not match:
            continue
        level = min(len(match.group(1)), 3)
        if level not in {2, 3}:
            continue
        original = clean_inline(match.group(2))
        if original == "Spis treści":
            titles.append((level, original, original))
            continue
        if level == 2:
            h2 += 1
            h3 = 0
            display = f"{h2}. {original}"
        else:
            if h2 == 0:
                display = original
            else:
                h3 += 1
                display = f"{h2}.{h3}. {original}"
        titles.append((level, original, display))
    return titles


def add_heading_bookmark(paragraph, name: str, bm_id: int) -> None:
    """Wrap a heading paragraph in a Word bookmark so links can target it."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_toc_entry(document: Document, anchor: str, text: str, level: int) -> None:
    """Add a clickable internal hyperlink that jumps to the heading bookmark."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.left_indent = Cm(0.55 if level == 3 else 0)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "23")  # half-points -> 11.5 pt
    rpr.append(size)
    run.append(rpr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text + " "
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    paragraph.add_run("." * max(8, 68 - len(text)))
    paragraph.add_run(" ")
    add_field(paragraph, f"PAGEREF {anchor} \\h")


def add_reference_entry(document: Document, label: str, anchor: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1
    add_inline_runs(paragraph, label)
    paragraph.add_run(" " + "." * max(8, 82 - len(label)))
    paragraph.add_run(" ")
    add_field(paragraph, f"PAGEREF {anchor} \\h")


def add_material_lists(document: Document, report_items: dict[str, list[dict[str, str | int]]]) -> None:
    intro = document.add_paragraph()
    intro.add_run(
        "Poniższe zestawienie zostało wygenerowane z treści raportu. Numery stron są polami referencyjnymi DOCX/PDF, tak jak w głównym spisie treści."
    )

    for title, key, prefix in [
        ("Spis rysunków", "figures", "Rysunek"),
        ("Spis tabel", "tables", "Tabela"),
        ("Spis listingów", "listings", "Listing"),
    ]:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string("1F4E79")
        for item in report_items[key]:
            add_reference_entry(
                document,
                f"{prefix} {item['number']}. {clean_inline(str(item['title']))}",
                str(item["anchor"]),
            )


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:r")
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin.append(begin_char)

    instr_run = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr_run.append(instr_text)

    separate = OxmlElement("w:r")
    separate_char = OxmlElement("w:fldChar")
    separate_char.set(qn("w:fldCharType"), "separate")
    separate.append(separate_char)

    end = OxmlElement("w:r")
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end.append(end_char)

    paragraph._p.append(begin)
    paragraph._p.append(instr_run)
    paragraph._p.append(separate)
    paragraph._p.append(end)


def add_word_toc(document: Document) -> None:
    note = document.add_paragraph()
    note.add_run(
        "Spis treści jest polem automatycznym Word/LibreOffice; po ręcznej edycji dokumentu należy użyć funkcji Aktualizuj pole."
    ).italic = True
    field_paragraph = document.add_paragraph()
    add_field(field_paragraph, r'TOC \o "1-3" \h \z \u')
    document.add_paragraph()


def add_page_number_footer(document: Document) -> None:
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = 1
    footer.add_run("Strona ")
    add_field(footer, "PAGE")


def convert_markdown(md_path: Path, out_path: Path) -> None:
    global ACTIVE_TOC_PAGE_NUMBERS, FIGURE_COUNTER, LISTING_COUNTER, TABLE_COUNTER
    ACTIVE_TOC_PAGE_NUMBERS = TOC_PAGE_NUMBERS if md_path.name == "Sprawozdanie_BAI.md" else {}
    FIGURE_COUNTER = 0
    LISTING_COUNTER = 0
    TABLE_COUNTER = 0

    output_dir = out_path.parent
    (output_dir / "assets").mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    section_titles = collect_section_titles(lines)
    report_items = collect_report_items(lines)
    index = 0
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    heading_counter = 0
    heading_occurrence = 0
    skip_manual_toc = False

    while index < len(lines):
        line = lines[index]

        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines, code_language)
                code_lines = []
                code_language = ""
                in_code = False
            else:
                in_code = True
                code_language = line.strip("`").strip()
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        # After the auto-generated clickable TOC, skip the manual TOC list in
        # the Markdown until the next section heading starts.
        if skip_manual_toc:
            terminator = re.match(r"^(#{1,6})\s+(.+)$", line)
            if terminator and len(terminator.group(1)) == 2:
                skip_manual_toc = False
            else:
                index += 1
                continue

        if not line.strip():
            index += 1
            continue

        if "page-break-after" in line or "page-break-before" in line:
            # The Markdown places a hard page break before every section, which
            # leaves dozens of half-empty pages in the DOCX/PDF. Let the body
            # flow naturally instead; only the title page and the table of
            # contents are isolated (handled in the heading branch below).
            index += 1
            continue

        image_match = IMAGE_RE.search(line)
        if image_match:
            add_image(document, md_path, output_dir, image_match.group(1), image_match.group(2))
            index += 1
            continue

        if line.lstrip().startswith("|"):
            table_rows, index = parse_table(lines, index)
            add_table(document, table_rows)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            text = clean_inline(heading_match.group(2))
            display_text = text
            if level in {2, 3}:
                if heading_occurrence < len(section_titles):
                    _, _, display_text = section_titles[heading_occurrence]
                heading_occurrence += 1
            paragraph = document.add_heading(display_text, level=level)
            # Bookmark every section heading so the table of contents can link
            # to it. Order matches collect_section_titles exactly.
            if level in {2, 3}:
                heading_counter += 1
                add_heading_bookmark(paragraph, f"sec{heading_counter}", heading_counter)
            # Front matter only: the table of contents and the first body
            # section each start on a fresh page. Everything else flows.
            if text in {"Spis treści", "Spis materiałów dokumentacyjnych", "Streszczenie"} or text.startswith("0. "):
                paragraph.paragraph_format.page_break_before = True
            # Replace the manual TOC list with a generated clickable one.
            if text == "Spis treści":
                toc_number = 0
                for level_number, original_title, display_title in section_titles:
                    toc_number += 1
                    if original_title == "Spis treści":
                        continue
                    add_toc_entry(document, f"sec{toc_number}", display_title, level_number)
                skip_manual_toc = True
            if text == "Spis materiałów dokumentacyjnych":
                add_material_lists(document, report_items)
            index += 1
            continue

        list_match = re.match(r"^\s*[-*]\s+(.+)$", line)
        if list_match:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, list_match.group(1))
            index += 1
            continue

        numbered_match = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
        if numbered_match:
            # Word's built-in List Number style tends to continue numbering
            # across unrelated Markdown lists after DOCX/PDF export. Keep the
            # author-provided Markdown number as plain text so each scenario
            # starts exactly where the report says it starts.
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.55)
            paragraph.paragraph_format.first_line_indent = Cm(-0.35)
            run = paragraph.add_run(f"{numbered_match.group(1)}. ")
            run.bold = False
            add_inline_runs(paragraph, numbered_match.group(2))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, line)
        index += 1

    if in_code and code_lines:
        add_code_block(document, code_lines, code_language)

    document.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    convert_markdown(args.markdown.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
